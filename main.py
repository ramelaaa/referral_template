from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create a new Document
doc = Document()

# Add a paragraph with custom font and size
paragraph = doc.add_paragraph()
run = paragraph.add_run('')
font = run.font
font.name = 'Times New Roman'  # Set the font name
font.size = Pt(14)  # Set the font size to 14 points


# Add centered header
header = doc.add_paragraph()
header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
header.add_run('Baychester Dental Multispecialties').bold = True
header.add_run('\n4000 Baychester Ave').bold = True
header.add_run('\nBronx, NY 10466').bold = True
header.add_run('\n(718) 618-6787').bold = True

# Add a blank line
doc.add_paragraph()

# Add "RE: " section aligned to the left and prompt for user name
re_paragraph = doc.add_paragraph()
re_paragraph.add_run('RE: ').bold = True
user_name = input('Enter your name: ')
re_paragraph.add_run(user_name)

# Add a blank line
doc.add_paragraph()

# Add "DOB: " section aligned to the left and prompt for date of birth
dob_paragraph = doc.add_paragraph()
dob_paragraph.add_run('DOB: ').bold = True
user_dob = input('Enter your date of birth (MM-DD-YYYY): ')
dob_paragraph.add_run(user_dob)

# Add a blank line
doc.add_paragraph()

# Add "Dear Dentist, " section aligned to the left
dear_paragraph = doc.add_paragraph()
dear_paragraph.add_run('Dear Dentist, ')

# Add a blank line
doc.add_paragraph()

# Prompt the user for specialty and list of numbers
specialty = input('Enter the type of specialty: ')
numbers = input('Enter a list of numbers (comma-separated): ')

# Add the referral information
referral_paragraph = doc.add_paragraph()
referral_paragraph.add_run('We are referring our patient ')
referral_paragraph.add_run(user_name).bold = True
referral_paragraph.add_run(' an ')
referral_paragraph.add_run(specialty).bold = True
referral_paragraph.add_run(', ')
referral_paragraph.add_run('for evaluation and treatment of #')
referral_paragraph.add_run(numbers).bold = True
referral_paragraph.add_run('. ')

# Add a blank line
doc.add_paragraph()

# Add some text to the document
thank_you_paragraph = doc.add_paragraph('Thank you!')
thank_you_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
doc.add_paragraph()
doc.add_paragraph()
dr_paragraph = doc.add_paragraph()
dr_paragraph.add_run('Dr. ').bold = True
doctor_name = input('Enter the last name of the doctor: ')
dr_paragraph.add_run(doctor_name).bold = True
dr_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

# Indent the referral paragraph
#referral_paragraph.paragraph_format.left_indent = Inches(0.5)

# Save the document
doc.save(user_name+'_referral_letter.docx')
